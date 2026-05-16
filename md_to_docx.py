"""Convert OFRN_in_Schools.md into a Word (.docx) document styled to match
the OFRN / UNL Extension publication look (red bold headings with red
horizontal rules, red-banded tables, branded footer).
"""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Inches, RGBColor

HERE = Path(__file__).parent
SRC_MD = HERE / "OFRN_in_Schools.md"
OUT_DOCX = HERE / "OFRN_in_Schools.docx"

# OFRN/UNL Extension brand color (deep red)
BRAND_RED = RGBColor(0xC0, 0x00, 0x00)
BRAND_RED_HEX = "C00000"

INLINE_RE = re.compile(
    r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`|\[[^\]]+\]\([^)]+\))'
)


# ---------- low-level XML helpers ------------------------------------------

def set_cell_shading(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def add_bottom_border(paragraph, hex_color: str = BRAND_RED_HEX,
                      size: str = "12") -> None:
    """Add a colored horizontal rule below a paragraph (for headings)."""
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)        # 1/8 pt units
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), hex_color)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def set_table_borders(table, hex_color: str = "000000",
                      size: str = "4") -> None:
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), size)
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), hex_color)
        borders.append(b)
    tbl_pr.append(borders)


def add_page_number_field(paragraph) -> None:
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr)
    run._r.append(fld_char2)
    run.font.size = Pt(9)
    run.font.name = "Calibri"


# ---------- inline parsing --------------------------------------------------

def add_inline(paragraph, text: str, base_color: RGBColor | None = None) -> None:
    for part in INLINE_RE.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2]); run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1]); run.italic = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1]); run.font.name = "Consolas"
        elif part.startswith("[") and "](" in part:
            label = part[1:part.index("]")]
            run = paragraph.add_run(label)
        else:
            run = paragraph.add_run(part)
        if base_color is not None:
            run.font.color.rgb = base_color


# ---------- block builders --------------------------------------------------

HEADING_SPECS = {
    1: {"size": 22, "rule": True,  "rule_size": "18"},
    2: {"size": 16, "rule": True,  "rule_size": "12"},
    3: {"size": 13, "rule": False},
    4: {"size": 12, "rule": False},
}


def add_heading(doc: Document, text: str, level: int) -> None:
    spec = HEADING_SPECS.get(level, HEADING_SPECS[3])
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level <= 2 else 8)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(spec["size"])
    run.font.color.rgb = BRAND_RED
    if spec.get("rule"):
        add_bottom_border(p, BRAND_RED_HEX, spec.get("rule_size", "12"))


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.autofit = True
    set_table_borders(table, "000000", "6")
    for r, row in enumerate(rows):
        for c, cell_text in enumerate(row):
            cell = table.rows[r].cells[c]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell.text = ""
            p = cell.paragraphs[0]
            if r == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                set_cell_shading(cell, BRAND_RED_HEX)
                add_inline(p, cell_text.strip(), base_color=RGBColor(0xFF, 0xFF, 0xFF))
                for run in p.runs:
                    run.bold = True
                    run.font.size = Pt(11)
            else:
                add_inline(p, cell_text.strip())
                for run in p.runs:
                    run.font.size = Pt(10)


# ---------- main parser -----------------------------------------------------

def parse_markdown(md_path: Path, doc: Document) -> None:
    lines = md_path.read_text(encoding="utf-8").splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped or stripped == "---":
            i += 1
            continue

        # Headings
        m = re.match(r'^(#{1,4})\s+(.*)$', stripped)
        if m:
            add_heading(doc, m.group(2), len(m.group(1)))
            i += 1
            continue

        # Pipe tables
        if stripped.startswith("|") and stripped.endswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip()); i += 1
            rows = []
            for tl in table_lines:
                if re.match(r'^\|[\s\-:|]+\|$', tl):
                    continue
                rows.append([c.strip() for c in tl.strip("|").split("|")])
            add_table(doc, rows); doc.add_paragraph()
            continue

        # Bullet lists
        if re.match(r'^(\s*)[-*] ', line):
            while i < len(lines) and re.match(r'^(\s*)[-*] ', lines[i]):
                m = re.match(r'^(\s*)[-*] (.*)$', lines[i])
                indent = len(m.group(1)) // 2
                style = "List Bullet" if indent == 0 else "List Bullet 2"
                p = doc.add_paragraph(style=style)
                add_inline(p, m.group(2)); i += 1
            continue

        # Numbered lists
        if re.match(r'^\s*\d+\.\s', line):
            while i < len(lines) and re.match(r'^\s*\d+\.\s', lines[i]):
                m = re.match(r'^\s*\d+\.\s+(.*)$', lines[i])
                p = doc.add_paragraph(style="List Number")
                add_inline(p, m.group(1)); i += 1
            continue

        # Blockquote
        if stripped.startswith("> "):
            qlines = []
            while i < len(lines) and lines[i].strip().startswith("> "):
                qlines.append(lines[i].strip()[2:]); i += 1
            p = doc.add_paragraph(style="Intense Quote")
            add_inline(p, " ".join(qlines))
            continue

        # Paragraph
        para_lines = []
        while i < len(lines):
            ln = lines[i]; s = ln.strip()
            if (not s or s == "---" or s.startswith(("#", "|", "> "))
                or re.match(r'^(\s*)[-*] ', ln)
                or re.match(r'^\s*\d+\.\s', ln)):
                break
            para_lines.append(s); i += 1
        if para_lines:
            p = doc.add_paragraph()
            add_inline(p, " ".join(para_lines))
        else:
            i += 1


# ---------- footer ----------------------------------------------------------

def build_footer(doc: Document) -> None:
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Top border on the footer paragraph for a separator line
    add_bottom_border(p, "808080", "4")  # using bottom border spec on top? use real top
    # Actually add a top border instead
    p_pr = p._p.get_or_add_pPr()
    # Remove the bottom border we just added, then add a top border
    for existing in p_pr.findall(qn("w:pBdr")):
        p_pr.remove(existing)
    pbdr = OxmlElement("w:pBdr")
    top = OxmlElement("w:top")
    top.set(qn("w:val"), "single")
    top.set(qn("w:sz"), "4")
    top.set(qn("w:space"), "4")
    top.set(qn("w:color"), "808080")
    pbdr.append(top)
    p_pr.append(pbdr)

    def add_run(text: str, *, bold=False) -> None:
        r = p.add_run(text)
        r.font.name = "Calibri"; r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0x59, 0x59, 0x59)
        r.bold = bold

    add_run("Nebraska On-Farm Research Network")
    add_run("  |  ")
    add_run("University of Nebraska–Lincoln")
    add_run("  |  Page ")
    add_page_number_field(p)


# ---------- main ------------------------------------------------------------

def main() -> None:
    doc = Document()

    # Body defaults
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    for section in doc.sections:
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)

    parse_markdown(SRC_MD, doc)
    build_footer(doc)

    doc.save(OUT_DOCX)
    print(f"Wrote {OUT_DOCX.name} ({OUT_DOCX.stat().st_size:,} bytes)")


if __name__ == "__main__":
    main()
